import io
import json
import os
import re

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, redirect, render_template, request, send_file, session, url_for
from dotenv import load_dotenv
from pypdf import PdfReader

load_dotenv()

app = Flask(__name__)
app.secret_key = os.urandom(24)

SYSTEM_PROMPT = """You are a contract analysis assistant. Analyse the provided contract text and return ONLY a valid JSON object — no markdown, no code fences, no prose, just raw JSON.

The JSON must conform to this exact schema:
{
  "parties": ["string"],
  "governing_law": "string",
  "term_months": number or null,
  "key_dates": [{"label": "string", "date": "string"}],
  "obligations": ["string"],
  "risk_clauses": [{"type": "string", "severity": "high" | "medium" | "low", "text": "string"}],
  "summary": "string (exactly 3 sentences in plain English)"
}

Rules:
- parties: list every named company or individual party
- governing_law: the jurisdiction governing the contract (e.g. "England and Wales")
- term_months: duration in months as a number, or null if not specified
- key_dates: important dates such as effective date, expiry, renewal deadlines
- obligations: concise plain-English statements of the key obligations on each party
- risk_clauses: identify clauses that pose legal, financial or operational risk; severity must be exactly "high", "medium" or "low"
- summary: exactly 3 sentences summarising what the contract is, who is involved, and what the main obligations/risks are

Return only the JSON object. Nothing else."""


def extract_text_from_pdf(file_stream):
    reader = PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.strip()


def analyse_contract(text):
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": f"Analyse this contract:\n\n{text}"
            }
        ]
    )
    raw = message.content[0].text.strip()
    # Strip markdown fences if Claude wraps anyway
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def build_docx(data):
    doc = Document()

    # Title
    title = doc.add_heading("Contract Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Parties
    doc.add_heading("Parties", level=1)
    for p in data.get("parties", []):
        doc.add_paragraph(p, style="List Bullet")

    # Governing Law
    doc.add_heading("Governing Law", level=1)
    doc.add_paragraph(data.get("governing_law", "Not specified"))

    # Term
    doc.add_heading("Contract Term", level=1)
    term = data.get("term_months")
    doc.add_paragraph(f"{term} months" if term else "Not specified")

    # Key Dates
    doc.add_heading("Key Dates", level=1)
    for kd in data.get("key_dates", []):
        doc.add_paragraph(f"{kd['label']}: {kd['date']}", style="List Bullet")

    # Obligations
    doc.add_heading("Key Obligations", level=1)
    for ob in data.get("obligations", []):
        doc.add_paragraph(ob, style="List Bullet")

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(data.get("summary", ""))

    # Risk Clauses
    doc.add_heading("Risk Clauses", level=1)
    severity_labels = {"high": "HIGH", "medium": "MEDIUM", "low": "LOW"}
    for rc in data.get("risk_clauses", []):
        severity = rc.get("severity", "low")
        p = doc.add_paragraph()
        run_label = p.add_run(f"[{severity_labels.get(severity, severity.upper())}] {rc.get('type', '')}: ")
        run_label.bold = True
        p.add_run(rc.get("text", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/analyse", methods=["POST"])
def analyse():
    if "contract" not in request.files or request.files["contract"].filename == "":
        return render_template("index.html", error="Please select a PDF file.")

    file = request.files["contract"]
    if not file.filename.lower().endswith(".pdf"):
        return render_template("index.html", error="Only PDF files are supported.")

    try:
        text = extract_text_from_pdf(file.stream)
        if not text:
            return render_template("index.html", error="Could not extract text from the PDF. Ensure it is not a scanned image.")
        data = analyse_contract(text)
        session["analysis"] = data
        return redirect(url_for("results"))
    except json.JSONDecodeError:
        return render_template("index.html", error="Claude returned an unexpected response. Please try again.")
    except anthropic.APIError as e:
        return render_template("index.html", error=f"API error: {e}")
    except Exception as e:
        return render_template("index.html", error=f"An error occurred: {e}")


@app.route("/results")
def results():
    data = session.get("analysis")
    if not data:
        return redirect(url_for("index"))
    return render_template("results.html", data=data)


@app.route("/download")
def download():
    data = session.get("analysis")
    if not data:
        return redirect(url_for("index"))
    buf = build_docx(data)
    return send_file(
        buf,
        as_attachment=True,
        download_name="contract_analysis.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    app.run(debug=True)
